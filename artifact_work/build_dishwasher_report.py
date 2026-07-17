from __future__ import annotations

import math
import os
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"D:\develop\food-swift")
WORK = WORKSPACE / "artifact_work"
OUTPUT = WORKSPACE / "第一次平时作业-现代洗碗机的历史与科伦设计演进分析.docx"
DIAGRAM = WORK / "cochrane_dishwasher_design.png"
KAI_FONT_FILE = Path(r"C:\Windows\Fonts\simkai.ttf")

BLACK = "000000"
FONT_NAME = "楷体"


TITLE = "第一次平时作业 - 现代洗碗机的历史与科伦设计演进分析"
META = "课程：工程与思维创新　｜　主题：技术史、系统设计与迭代修正"

ABSTRACT = (
    "现代洗碗机并非一次完成的“灵感产品”，而是从手摇拨水、餐具约束、压力喷淋到过滤循环和分阶段漂洗逐步形成的系统。"
    "本文以约瑟芬·科伦（Josephine G. Cochran/Cochrane）的专利链为主线，分析她如何把“少破损、洗得净、洗得快”转译为可制造、可维护的工程结构，并说明这些结构怎样延续到今天。"
)

HISTORY_P1 = (
    "机械洗碗的尝试早于科伦。1850年，乔尔·霍顿的专利把餐具放入筒体，以手柄带动带叶片的轴，把水泼向餐具；1865年，利维·亚历山大又改进餐具架、挡水边和排水阀。"
    "二者证明“机器可以替人搅动水”，却没有同时解决餐具固定、冲洗覆盖和可靠循环，因此没有形成实用产品。[1][2]"
)

HISTORY_P2 = (
    "转折发生在19世纪80年代。科伦因珍贵瓷器在手洗中被磕坏而产生设计动机，1883年丧偶与债务又迫使她把想法变成商品。"
    "她与机械师乔治·巴特斯制作样机，1885年末提交申请，1886年获美国专利355,139。她并非“第一个想到洗碗机的人”，但她首次把按餐具形状制作的金属丝笼、稳定的转架与压力水流组合起来，使机器能够反复、较少破损地完成洗涤。[3][8]"
)

HISTORY_P3 = (
    "1893年芝加哥哥伦布纪念博览会成为商业验证：展会餐厅实际使用多台机器，评委以机械结构、耐久性和适用性授予最高奖。"
    "早期客户主要是酒店、餐馆、医院和学校，因为它们有集中热水、足够空间和高周转需求；家庭当时既缺稳定热水，也缺合适的洗涤剂。1949年，基于其技术脉络的KitchenAid洗碗机进入公众市场，20世纪50年代以后才逐渐普及。[8][9]"
)

CORE_P1 = (
    "1886年方案的核心不是“用电”，而是对水、餐具和运动的组织。盘、杯、碟和刀叉被分别夹持在金属丝架中，餐具架装入可转动的笼式转架；泵把热肥皂水或清洁热水送入喷管，转架让不同表面依次进入水流。"
    "机体下部设两个相互分隔的液槽和两套泵，活动底板/导流件把回流水送回对应液槽，从而把洗涤与清洗介质分开并重复利用。[3]"
)

CORE_P2 = (
    "这套结构包含了现代洗碗机的四个基本模块：餐具定位、压力喷淋、分阶段水路和封闭循环。后来的家用机把“转动餐具”改为“固定餐具、旋转喷臂”，但其目的没有改变：用可控相对运动扩大覆盖率；今天的滤网、循环泵、洗涤—漂洗程序，也能在科伦的专利中找到明确的工程祖型。"
)

ITERATION_INTRO = (
    "科伦的价值还在于她没有把首件样机当成终点。后续专利反复针对堵塞、喷淋不均、转动间歇、换装低效和漂洗死角进行修正，形成了“发现失效—拆分功能—改结构—再验证”的连续迭代。"
)

ITERATION_ROWS = [
    (
        "1886\nUS355,139",
        "手洗易破损；早期机器靠浸泡、刷洗或随意泼水，覆盖不稳定。",
        "设置盘碟、杯具和刀叉专用金属丝架；转架与压力喷管配合；双液槽、双泵及活动导流底板分开热肥皂水与清洁热水。形成“固定—喷淋—循环—切换”的基础架构。[3]",
    ),
    (
        "1888\nUS391,782",
        "食物残渣会堵泵；长喷管前后出水不均；人工切换回流水路易出错。",
        "在泵前增加多孔过滤板；在长槽喷管内加入带孔辅管与反向导片，使全长均匀出水；泵杆联动摆动导流板，自动把回水送入正确液槽。[4]",
    ),
    (
        "1894\nUS512,683",
        "间歇旋转造成清洗盲区；停机装卸降低商业场景吞吐量；餐具可能松动。",
        "把转架改为连续旋转；轨道小车可同时承载内外两只转架，一只清洗、一只装卸；以齿条、弹簧卡爪和压紧件锁定餐具，并简化自动导流结构。[5]",
    ),
    (
        "1911\nUS1,009,223",
        "单方向水流难以触及凹面和背面；运动喷嘴可能碰伤餐具；开口机体易飞溅。",
        "在餐具上、下布置可旋转喷嘴，利用反作用力扩大角度；增加防碰保护、可调供水、分段联动机门和餐具架导轨。现代旋转喷臂与前门装载由此更清晰。[6]",
    ),
    (
        "1917\nUS1,223,380",
        "循环水要节约，但最终漂洗又必须卫生；固定喷射仍可能留下局部死角。",
        "形成全封闭洗/漂系统：肥皂水经过滤、除渣和除油后循环，外部新鲜热水只用于一次漂洗；操作者可从机外摆动漂洗管，并配合旋转喷头覆盖不同餐架。该专利为身后授权。[7]",
    ),
]

ANALYSIS_P1 = (
    "在我看来，科伦最重要的创新方法是把生活抱怨改写成工程指标。“别再打碎瓷器”对应约束与缓冲，“洗得干净”对应水压、角度和相对运动，“酒店能赚钱”对应连续装卸与高吞吐。"
    "因此，她的专利不是零件堆积，而是围绕用户风险建立功能链。"
)

ANALYSIS_P2 = (
    "第二，她善于处理矛盾：水流要强，却不能伤盘；水要循环节省，却不能把污物带回；机体要封闭防飞溅，又要方便装卸。"
    "1888年的过滤与自动导流、1894年的双转架、1911年的上下旋转喷嘴、1917年的“循环洗涤水＋一次性新鲜漂洗水”，本质上都是把冲突功能解耦。"
)

ANALYSIS_P3 = (
    "第三，技术能否成为“现代产品”取决于系统环境。科伦先进入酒店而非普通家庭，说明好发明也需要热水、电力、洗涤剂、制造成本与消费观念共同成熟。"
    "现代洗碗机加入电控、传感器、变速泵、加热干燥和节能程序，但底层逻辑仍是她确立的：可靠固定餐具，用受控水流覆盖表面，把洗、滤、漂、排组织成可重复流程。"
)

CONCLUSION = (
    "科伦的贡献不在于凭空创造“自动洗碗”，而在于把不可靠的机械想法推进为可验证、可销售、可持续修正的系统。她的历次改进说明：工程创新通常不是一次命中，而是让每次失败都暴露一个可被重新设计的接口。"
)

REFERENCES = [
    ("[1]", "Joel Houghton, US Patent 7,365, Improvement in machines for washing table furniture, 1850.", "https://patents.google.com/patent/US7365A/en"),
    ("[2]", "Levi A. Alexander, US Patent 51,000, Machine for washing dishes, 1865.", "https://patents.google.com/patent/US51000A/en"),
    ("[3]", "Josephine G. Cochran, US Patent 355,139, Dish-washing machine, 1886.", "https://patents.google.com/patent/US355139A/en"),
    ("[4]", "Josephine G. Cochran & Jacob Kritch, US Patent 391,782, Dish-washing machine, 1888.", "https://patents.google.com/patent/US391782A/en"),
    ("[5]", "Josephine G. Cochrane, US Patent 512,683, Dish-cleaner, 1894.", "https://patents.google.com/patent/US512683A/en"),
    ("[6]", "Josephine G. Cochrane, US Patent 1,009,223, Dish-washing machine, 1911.", "https://patents.google.com/patent/US1009223A/en"),
    ("[7]", "Josephine Garis-Cochrane, US Patent 1,223,380, Dish-washing machine, 1917.", "https://patents.google.com/patent/US1223380A/en"),
    ("[8]", "United States Patent and Trademark Office, “I’ll do it myself.”", "https://www.uspto.gov/learning-and-resources/journeys-innovation/historical-stories/ill-do-it-myself"),
    ("[9]", "National Inventors Hall of Fame, Josephine Garis Cochran.", "https://www.invent.org/inductees/josephine-garis-cochran"),
]


def font(size: int):
    return ImageFont.truetype(str(KAI_FONT_FILE), size=size)


IMG_TITLE = font(52)
IMG_SECTION = font(38)
IMG_LABEL = font(30)
IMG_SMALL = font(26)
IMG_TINY = font(23)


def text_bbox(draw: ImageDraw.ImageDraw, text: str, fnt) -> tuple[int, int, int, int]:
    return draw.textbbox((0, 0), text, font=fnt)


def draw_centered(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fnt, spacing: int = 7):
    x1, y1, x2, y2 = xy
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align="center")
    w, h = box[2] - box[0], box[3] - box[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=fnt, fill=0, spacing=spacing, align="center")


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> str:
    lines: list[str] = []
    current = ""
    for ch in text:
        if ch == "\n":
            lines.append(current)
            current = ""
            continue
        trial = current + ch
        if current and text_bbox(draw, trial, fnt)[2] > max_width:
            lines.append(current)
            current = ch
        else:
            current = trial
    if current:
        lines.append(current)
    return "\n".join(lines)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 5, head: int = 18):
    draw.line([start, end], fill=0, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    p1 = (end[0] - head * math.cos(ang - math.pi / 6), end[1] - head * math.sin(ang - math.pi / 6))
    p2 = (end[0] - head * math.cos(ang + math.pi / 6), end[1] - head * math.sin(ang + math.pi / 6))
    draw.polygon([end, p1, p2], fill=0)


def dashed_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 4, dash: int = 18, gap: int = 12):
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    ux, uy = (x2 - x1) / length, (y2 - y1) / length
    pos = 0.0
    while pos < max(0, length - 22):
        end_pos = min(pos + dash, length - 22)
        draw.line([(x1 + ux * pos, y1 + uy * pos), (x1 + ux * end_pos, y1 + uy * end_pos)], fill=0, width=width)
        pos += dash + gap
    arrow(draw, (int(x2 - ux * 26), int(y2 - uy * 26)), end, width=width, head=16)


def numbered_callout(draw: ImageDraw.ImageDraw, number: int, label: str, anchor: tuple[int, int], box: tuple[int, int, int, int]):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=15, outline=0, width=4, fill=255)
    circle = (x1 + 12, y1 + 12, x1 + 62, y1 + 62)
    draw.ellipse(circle, outline=0, width=4, fill=255)
    draw_centered(draw, circle, str(number), IMG_SMALL)
    wrapped = wrap_text(draw, label, IMG_SMALL, x2 - x1 - 88)
    draw.multiline_text((x1 + 76, y1 + 13), wrapped, font=IMG_SMALL, fill=0, spacing=5)
    edge = (x2, (y1 + y2) // 2) if anchor[0] > x2 else (x1, (y1 + y2) // 2)
    arrow(draw, edge, anchor, width=3, head=13)


def evolution_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], year: str, problem: str, fix: str):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=20, outline=0, width=5, fill=255)
    draw.line([(x1 + 170, y1), (x1 + 170, y2)], fill=0, width=3)
    draw_centered(draw, (x1 + 8, y1 + 8, x1 + 162, y2 - 8), year, IMG_SECTION)
    p = wrap_text(draw, "问题：" + problem, IMG_TINY, x2 - x1 - 200)
    f = wrap_text(draw, "修正：" + fix, IMG_TINY, x2 - x1 - 200)
    draw.multiline_text((x1 + 190, y1 + 16), p, font=IMG_TINY, fill=0, spacing=4)
    p_lines = p.count("\n") + 1
    draw.multiline_text((x1 + 190, y1 + 22 + p_lines * 31), f, font=IMG_TINY, fill=0, spacing=4)


def create_diagram(path: Path):
    img = Image.new("L", (3000, 1850), color=255)
    draw = ImageDraw.Draw(img)
    draw.rectangle((24, 24, 2976, 1826), outline=0, width=5)
    draw_centered(draw, (70, 32, 2930, 112), "科伦洗碗机：核心结构与历次修正设计图", IMG_TITLE)
    draw_centered(draw, (70, 110, 2930, 160), "依据1886、1888、1894、1911、1917年专利重绘；非等比例；虚线表示水流", IMG_SMALL)

    draw.text((92, 190), "A　结构剖视与工作回路", font=IMG_SECTION, fill=0)
    draw.text((1900, 190), "B　五次迭代修正链", font=IMG_SECTION, fill=0)

    # Main cabinet and wash chamber.
    outer = (500, 300, 1720, 1250)
    inner = (570, 365, 1650, 920)
    draw.rounded_rectangle(outer, radius=25, outline=0, width=7, fill=255)
    draw.rectangle(inner, outline=0, width=5)
    draw.line([(1650, 365), (1718, 320)], fill=0, width=4)
    draw.line([(1650, 920), (1718, 965)], fill=0, width=4)
    draw.text((1500, 385), "分段机门", font=IMG_TINY, fill=0)

    # Rack, plates and guide rail.
    draw.rectangle((710, 470, 1490, 805), outline=0, width=4)
    draw.line([(680, 825), (1530, 825)], fill=0, width=6)
    draw.line([(720, 805), (700, 825), (720, 845)], fill=0, width=4)
    draw.line([(1480, 805), (1500, 825), (1480, 845)], fill=0, width=4)
    for i, cx in enumerate([820, 950, 1080, 1210, 1340]):
        draw.ellipse((cx - 38, 530 - (i % 2) * 10, cx + 38, 725 + (i % 2) * 10), outline=0, width=4)
        draw.line([(cx - 58, 755), (cx, 720), (cx + 58, 755)], fill=0, width=3)
    draw.line([(650, 637), (1550, 637)], fill=0, width=5)
    draw.ellipse((620, 607, 680, 667), outline=0, width=4)
    draw.ellipse((1520, 607, 1580, 667), outline=0, width=4)

    # Drive input and rotation arrow.
    draw.ellipse((410, 560, 500, 650), outline=0, width=5)
    draw.line([(455, 605), (375, 520), (330, 555)], fill=0, width=6)
    draw.ellipse((312, 540, 342, 570), outline=0, width=4)
    draw.arc((725, 430, 1475, 820), 205, 330, fill=0, width=6)
    arrow(draw, (1380, 735), (1448, 650), width=6, head=22)

    # Upper and lower spray arms, holes and water jets.
    draw.line([(700, 420), (1500, 420)], fill=0, width=13)
    draw.line([(700, 865), (1500, 865)], fill=0, width=13)
    for cx in range(760, 1461, 140):
        draw.ellipse((cx - 8, 412, cx + 8, 428), fill=0)
        draw.ellipse((cx - 8, 857, cx + 8, 873), fill=0)
        dashed_arrow(draw, (cx, 430), (cx + (18 if cx % 280 == 0 else -18), 520), width=3)
        dashed_arrow(draw, (cx, 855), (cx + (-18 if cx % 280 == 0 else 18), 760), width=3)

    # Filter plate and return deflector.
    draw.line([(575, 940), (1645, 940)], fill=0, width=7)
    for x in range(600, 1641, 42):
        draw.ellipse((x, 933, x + 11, 944), fill=0)
    draw.polygon([(950, 950), (1110, 1020), (1270, 950)], outline=0)
    draw.line([(1110, 1020), (1110, 1090)], fill=0, width=5)

    # Dual reservoirs.
    draw.rectangle((575, 1025, 1100, 1185), outline=0, width=4)
    draw.rectangle((1120, 1025, 1645, 1185), outline=0, width=4)
    draw_centered(draw, (590, 1050, 1080, 1165), "洗涤液槽\n热肥皂水循环", IMG_LABEL)
    draw_centered(draw, (1140, 1050, 1625, 1165), "清水槽\n清洁热水切换", IMG_LABEL)

    # Pumps and supply risers.
    draw.ellipse((600, 1185, 760, 1345), outline=0, width=5)
    draw.ellipse((1455, 1185, 1615, 1345), outline=0, width=5)
    draw_centered(draw, (600, 1185, 760, 1345), "泵1", IMG_LABEL)
    draw_centered(draw, (1455, 1185, 1615, 1345), "泵2", IMG_LABEL)
    draw.line([(680, 1185), (680, 970), (620, 970), (620, 420), (700, 420)], fill=0, width=6)
    draw.line([(1535, 1185), (1535, 970), (1600, 970), (1600, 865), (1500, 865)], fill=0, width=6)
    arrow(draw, (620, 610), (620, 500), width=5, head=18)
    arrow(draw, (1600, 970), (1600, 890), width=5, head=18)

    # Fresh rinse input and drain.
    draw.line([(1650, 510), (1810, 510)], fill=0, width=6)
    arrow(draw, (1810, 510), (1660, 510), width=5, head=18)
    draw.text((1695, 465), "新鲜热水", font=IMG_TINY, fill=0)
    draw.line([(1110, 1185), (1110, 1345), (1260, 1345)], fill=0, width=6)
    arrow(draw, (1190, 1345), (1260, 1345), width=5, head=18)
    draw.text((1268, 1312), "溢流/排放", font=IMG_TINY, fill=0)

    # Callouts on the left.
    numbered_callout(draw, 1, "曲柄或动力输入\n带动转架", (455, 605), (80, 315, 430, 430))
    numbered_callout(draw, 2, "专用金属丝架\n分隔并夹持餐具", (760, 560), (80, 470, 430, 600))
    numbered_callout(draw, 3, "上、下喷管\n扩大冲洗角度", (720, 420), (80, 650, 430, 780))
    numbered_callout(draw, 4, "导轨与分段机门\n便于整架装卸", (710, 825), (80, 830, 430, 960))
    numbered_callout(draw, 5, "多孔滤板\n阻止残渣进入泵", (650, 940), (80, 1010, 430, 1140))
    numbered_callout(draw, 6, "摆动导流板\n把回水送入正确槽", (1110, 990), (80, 1190, 430, 1320))

    # Evolution boxes.
    evo_x1, evo_x2 = 1900, 2915
    evolutions = [
        ("1886", "餐具破损、覆盖不稳", "专用餐架＋压力喷淋＋双液槽"),
        ("1888", "残渣堵泵、喷管前后不均", "过滤板＋自动导流＋均匀出水"),
        ("1894", "旋转间歇、装卸停机", "连续转架＋双架换装＋弹簧压紧"),
        ("1911", "背面死角、飞溅与碰撞", "上下旋转喷嘴＋防护＋联动机门"),
        ("1917", "节水与卫生、局部死角冲突", "过滤回用洗涤水＋一次热水漂洗＋摆动喷管"),
    ]
    ys = [280, 520, 760, 1000, 1240]
    for idx, (year, problem, fix) in enumerate(evolutions):
        box = (evo_x1, ys[idx], evo_x2, ys[idx] + 190)
        evolution_box(draw, box, year, problem, fix)
        if idx < len(evolutions) - 1:
            arrow(draw, ((evo_x1 + evo_x2) // 2, ys[idx] + 190), ((evo_x1 + evo_x2) // 2, ys[idx + 1] - 12), width=5, head=18)

    # Bottom process flow.
    draw.line([(70, 1510), (2930, 1510)], fill=0, width=4)
    draw.text((92, 1540), "C　工作流程（结构逻辑）", font=IMG_SECTION, fill=0)
    flow = ["整架装载", "泵送洗涤液", "压力喷淋＋相对运动", "过滤并回流", "切换新鲜热水漂洗", "排放/取出"]
    x = 95
    y1, y2 = 1625, 1765
    widths = [360, 400, 560, 380, 550, 390]
    boxes = []
    for label, w in zip(flow, widths):
        box = (x, y1, x + w, y2)
        draw.rounded_rectangle(box, radius=18, outline=0, width=4, fill=255)
        draw_centered(draw, box, wrap_text(draw, label, IMG_SMALL, w - 30), IMG_SMALL)
        boxes.append(box)
        x += w + 35
    for a, b in zip(boxes, boxes[1:]):
        arrow(draw, (a[2] + 8, (a[1] + a[3]) // 2), (b[0] - 8, (b[1] + b[3]) // 2), width=4, head=16)

    img.save(path, format="PNG", dpi=(300, 300), optimize=True)


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None, underline: bool | None = None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_style_font(style, size: float, bold: bool = False):
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill="FFFFFF"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    trPr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")


def set_paragraph_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, 9)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9)


def style_footer(section):
    section.footer_distance = Inches(0.45)
    footer = section.footer
    footer.is_linked_to_previous = True
    p = footer.paragraphs[0]
    if not p.text and len(p._p) <= 2:
        add_page_number(p)


def apply_page_setup(section, landscape=False, figure_override=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    margin = 0.62 if figure_override else 1.0
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    style_footer(section)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(TITLE)
    set_run_font(r, 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(META)
    set_run_font(r, 10.5, False)


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 13, True)
    return p


def add_body(doc: Document, text: str, first_line=True, keep_with_next=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0.33) if first_line else Inches(0)
    p.paragraph_format.keep_with_next = keep_with_next
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def add_labeled_paragraph(doc: Document, label: str, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(label)
    set_run_font(r, 11, True)
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def add_picture_with_alt(paragraph, image_path: Path, width_inches: float, alt_text: str):
    run = paragraph.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width_inches))
    docPr = inline._inline.docPr
    docPr.set("descr", alt_text)
    docPr.set("title", "科伦洗碗机结构与改进设计图")


def add_evolution_table(doc: Document):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True
    r = caption.add_run("表1　科伦洗碗机专利节点中的问题与修正")
    set_run_font(r, 10.5, True)

    table = doc.add_table(rows=1, cols=3)
    headers = ["节点", "暴露的问题", "变更与修正"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_shading(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        set_run_font(r, 10, True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])

    for node, issue, change in ITERATION_ROWS:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, text in enumerate((node, issue, change)):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=110, bottom=110)
            set_cell_shading(cell, "FFFFFF")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run_font(r, 9.5, bold=(idx == 0))

    set_table_geometry(table, [1350, 2700, 5310], indent_dxa=120)
    table.style = "Table Grid"
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    r = after.add_run("说明：1917年专利在科伦去世后由遗产执行人办理授权，故文中将其视为其最终改进工作的延续。")
    set_run_font(r, 9)


def add_reference(doc: Document, marker: str, title: str, url: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(f"{marker} {title} ")
    set_run_font(r, 9)
    r = p.add_run(url)
    set_run_font(r, 8.5, underline=True)


def add_section_break(doc: Document, landscape=False, figure_override=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = True
    apply_page_setup(section, landscape=landscape, figure_override=figure_override)
    return section


def create_document(output_path: Path):
    doc = Document()
    configure_styles(doc)
    apply_page_setup(doc.sections[0], landscape=False, figure_override=False)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "现代洗碗机历史与约瑟芬·科伦专利迭代分析"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "洗碗机, 约瑟芬·科伦, 工程创新, 专利史"

    add_title_block(doc)
    add_labeled_paragraph(doc, "摘要：", ABSTRACT)
    add_heading(doc, "一、从“搅动水”到可用的洗碗系统", 1)
    add_body(doc, HISTORY_P1)
    add_body(doc, HISTORY_P2)
    add_body(doc, HISTORY_P3)

    add_heading(doc, "二、1886年核心设计：把餐具、水路与运动组织起来", 1)
    add_body(doc, CORE_P1)
    add_body(doc, CORE_P2)

    # A dedicated landscape page makes the technical drawing legible.
    add_section_break(doc, landscape=True, figure_override=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("图1　科伦洗碗机核心结构与历次修正（分析性重绘）")
    set_run_font(r, 13, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_picture_with_alt(
        p,
        DIAGRAM,
        9.55,
        "黑白设计图：左侧为科伦洗碗机核心结构剖视和六个部件标注，右侧为1886至1917五次专利修正链，底部为工作流程。",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("图中把多个年代的关键结构叠加在同一剖面中，目的在于展示功能演进，而非复原某一台具体机器。")
    set_run_font(r, 9)

    add_section_break(doc, landscape=False, figure_override=False)
    add_heading(doc, "三、历次变更和修正：从“能洗”到“稳定、高效、卫生”", 1)
    add_body(doc, ITERATION_INTRO)
    add_evolution_table(doc)

    add_heading(doc, "四、工程与思维创新分析", 1)
    add_body(doc, ANALYSIS_P1)
    add_body(doc, ANALYSIS_P2)
    add_body(doc, ANALYSIS_P3)

    add_heading(doc, "五、结论", 1)
    add_body(doc, CONCLUSION)

    add_heading(doc, "参考资料", 1)
    for marker, title, url in REFERENCES:
        add_reference(doc, marker, title, url)

    # Final global pass: keep every textual run in KaiTi and pure black.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, run.font.size.pt if run.font.size else None, run.bold, run.italic, run.underline)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, run.font.size.pt if run.font.size else None, run.bold, run.italic, run.underline)

    doc.save(output_path)


def count_chinese(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def main():
    WORK.mkdir(parents=True, exist_ok=True)
    create_diagram(DIAGRAM)
    create_document(OUTPUT)
    report_text = "".join(
        [
            ABSTRACT,
            HISTORY_P1,
            HISTORY_P2,
            HISTORY_P3,
            CORE_P1,
            CORE_P2,
            ITERATION_INTRO,
            *("".join(row) for row in ITERATION_ROWS),
            ANALYSIS_P1,
            ANALYSIS_P2,
            ANALYSIS_P3,
            CONCLUSION,
        ]
    )
    print(f"Created: {OUTPUT}")
    print(f"Diagram: {DIAGRAM}")
    print(f"Main-content Chinese character count: {count_chinese(report_text)}")


if __name__ == "__main__":
    main()
